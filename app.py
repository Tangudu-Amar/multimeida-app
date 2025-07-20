from flask import Flask, render_template, request, url_for
import google.generativeai as genai
from PIL import Image
import logging
import cv2
import io
import os
import requests
import whisper
import moviepy.editor as mp
from werkzeug.utils import secure_filename
import time
import google.api_core.exceptions
from docx import Document
from flask import send_file,session,after_this_request
from dotenv import load_dotenv
import mysql.connector
from mysql.connector import Error, InterfaceError
from flask import redirect, session, flash
from werkzeug.security import generate_password_hash, check_password_hash
import re  # Add this at the top of your file if not already


app = Flask(__name__)
app.secret_key = os.getenv("SECRET_KEY")
# Setup logging
os.makedirs("logs", exist_ok=True)
logging.basicConfig(
    filename='logs/app.log',  # Save logs to this file
    level=logging.INFO,        # You can also use DEBUG, WARNING, ERROR
    format='%(asctime)s %(levelname)s: %(message)s [in %(pathname)s:%(lineno)d]'
)

load_dotenv()
app.config['UPLOAD_FOLDER'] = 'static/uploads'
app.config['MAX_CONTENT_LENGTH'] = 20 * 1024 * 1024  # 20 MB limit

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Configure Gemini API key
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

# Load Whisper model for audio transcription
whisper_model = whisper.load_model("base")

    
def get_db_connection():
    try:
        connection = mysql.connector.connect(
            host=os.getenv("DB_HOST", "localhost"),
            user=os.getenv("DB_USER", "root"),
            password=os.getenv("DB_PASSWORD", ""),
            database=os.getenv("DB_NAME", "multimedia_app")
        )
        if connection.is_connected():
            print("✅ Successfully connected to the database.")
            return connection
    except InterfaceError as ie:
        print("❌ Could not connect to MySQL server.")
        print(f"InterfaceError: {ie}")
    except Error as e:
        print("❌ MySQL error occurred.")
        print(f"MySQLError: {e}")
    return None


@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']

        conn = get_db_connection()
        if conn is None:
            flash("⚠ Cannot connect to the database. Try again later.")
            return redirect(url_for('login'))

        try:
            cursor = conn.cursor(dictionary=True)
            cursor.execute("SELECT * FROM users WHERE username = %s", (username,))
            user = cursor.fetchone()
            conn.close()

            if user and check_password_hash(user['password'], password):
                session['username'] = username
                return redirect(url_for('index'))
            elif not user:
                flash("User doesn't exist. Please sign up.")
                return redirect(url_for('signup'))
            else:
                flash("Incorrect password.")
                return redirect(url_for('login'))

        except Exception as e:
            print("❌ Login DB error:", e)
            flash("⚠ Something went wrong during login.")
            return redirect(url_for('login'))

    return render_template('login.html')


@app.route("/signup", methods=["GET", "POST"])
def signup():
    if request.method == "POST":
        username = request.form["username"]
        raw_password = request.form["password"]

        # ✅ Password validation (8–16 chars, 1 uppercase, 1 lowercase, 1 number, 1 symbol)
        pattern = r'^(?=.*[a-z])(?=.*[A-Z])(?=.*\d)(?=.*[@$!%*?&])[A-Za-z\d@$!%*?&]{8,16}$'
        if not re.match(pattern, raw_password):
            flash("❌ Password must be 8–16 characters, include 1 number, 1 uppercase, 1 lowercase, and 1 special symbol.")
            return redirect(url_for("signup"))

        # ✅ Only hash the password after validation
        password = generate_password_hash(raw_password)

        conn = get_db_connection()
        if conn is None:
            flash("⚠ Cannot connect to the database. Try again later.")
            return redirect(url_for('signup'))

        try:
            cursor = conn.cursor()
            cursor.execute("SELECT id FROM users WHERE username = %s", (username,))
            existing_user = cursor.fetchone()

            if existing_user:
                flash("⚠ Username already exists.")
                return redirect(url_for("signup"))

            cursor.execute("INSERT INTO users (username, password) VALUES (%s, %s)", (username, password))
            conn.commit()
            conn.close()

            flash("✅ Signup successful. You can now log in.")
            return redirect(url_for("login"))

        except Exception as e:
            print("❌ Signup DB error:", e)
            flash("⚠ An error occurred during signup.")
            return redirect(url_for("signup"))

    return render_template("signup.html")


@app.route('/')
def index():
    if 'username' in session:
        return render_template('index.html')
    return redirect(url_for('login'))

@app.route('/logout', methods=['POST'])
def logout():
    session.pop('username', None)
    flash("✅ Logged out successfully.")
    return redirect(url_for('login'))


@app.route("/analyze", methods=["POST"])
def analyze():
    logging.info("🔍 Received /analyze request")
    file_type = request.form.get("file_type")

    if file_type == "image":
        return analyze_image()
    elif file_type == "audio":
        return analyze_audio()
    elif file_type == "video":
        return analyse_video()

    return render_template("index.html", summary="❌ Invalid file type."), 400

def analyze_image():
    file = request.files.get("image")
    image_url = request.form.get("image_url")
    image_bytes = None
    image_display_path = None

    if file and file.filename != "":
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        with open(file_path, "rb") as f:
            image_bytes = f.read()
        image_display_path = url_for('static', filename=f'uploads/{filename}')

    elif image_url and image_url.strip() != "":
        try:
            response = requests.get(image_url, timeout=10)
            response.raise_for_status()
            content_type = response.headers.get("Content-Type", "")
            if not content_type.startswith("image/"):
                return render_template("index.html", summary="❌ URL does not point to an image."), 400
            image_bytes = response.content
            filename = os.path.basename(image_url.split("?")[0])
            if not any(filename.lower().endswith(ext) for ext in [".jpg", ".jpeg", ".png", ".gif", ".webp"]):
                filename += ".jpg"
            filename = secure_filename(filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            with open(file_path, "wb") as f:
                f.write(image_bytes)
            image_display_path = url_for('static', filename=f'uploads/{filename}')
        except Exception as e:
            return render_template("index.html", summary=f"❌ Failed to fetch image: {str(e)}"), 400
    else:
        return render_template("index.html", summary="❌ No image or URL provided."), 400

    try:
        image = Image.open(io.BytesIO(image_bytes))
    except Exception as e:
        return render_template("index.html", summary=f"❌ Unable to process image: {str(e)}"), 400

    type_prompt = """
        You are a content classification expert. Given an image, classify it into **one** of the following exact categories only:

        People, Scenery, Graphs, UI, Text, ID, Product, Animals, Objects, Others

        Instructions:
        - Return just **one** of the above words.
        - Do NOT use any symbols, markdown, or explanation.
        - If the image doesn’t clearly fall into any category, return: Others
       """
    model = genai.GenerativeModel("models/gemini-2.0-flash")
    imgtype_response = model.generate_content([image, type_prompt])
    imgtype = imgtype_response.text.strip().lower()

    custom_prompt = """
    You are an expert visual analyst. Describe the image as clearly and objectively as possible. Focus on:
    - Key objects or subjects
    - Any notable visual features
    - Try to infer the likely purpose, context, or mood of the image
    Avoid formatting (no markdown, no symbols). Keep it clean and structured.
    """

    if imgtype == "people":
        custom_prompt += """
        This image shows people or portraits. Describe:
        - Number of individuals, gender, and age estimates
        - Clothing and facial expressions
        - Activity, posture, and background setting
        - Possible social setting or context
        """

    elif imgtype == "scenery":
        custom_prompt += """
        This image shows scenery or a landscape. Describe:
        - Natural or urban environment
        - Weather, lighting, time of day
        - Key objects (trees, buildings, roads, etc.)
        - Emotional or aesthetic tone
        """

    elif imgtype == "graphs":
        custom_prompt += """
        This image contains a graph or chart. Describe:
        - Type of graph (bar, line, pie, etc.)
        - Data trends or key observations
        - Any visible axes or labels
        - Likely domain (finance, health, education, etc.)
        """

    elif imgtype == "ui":
        custom_prompt += """
        This is a user interface screenshot. Describe:
        - Type of UI (website, app, dashboard)
        - Visible UI elements (buttons, tables, inputs, charts)
        - Branding, color scheme, or theme
        - Inferred purpose or user goal
        """

    elif imgtype == "text":
        custom_prompt += """
        This image shows text content. Describe:
        - Main headings or paragraph structure
        - Language and tone
        - Formatting (font size, bold/italic, layout)
        - What the text is trying to communicate
        """

    elif imgtype == "id":
        custom_prompt += """
        This image shows an ID or document. Describe:
        - Type of document (passport, Aadhar, license, etc.)
        - Layout: photo placement, fields (name, ID, DOB)
        - Any security features, logos, or text
        - Purpose of the document
        """

    elif imgtype == "product":
        custom_prompt += """
        This image shows a product. Describe:
        - Product type and appearance
        - Brand logos or labels
        - Background setup (studio, plain, natural)
        - Inferred use-case or target customer
        """

    elif imgtype == "animals":
        custom_prompt += """
        This image shows animals. Describe:
        - Type of animal(s) (dog, cat, bird, etc.)
        - Number, appearance, and actions
        - Environment (indoor, outdoor, zoo, wild)
        - Mood or tone of the image
        """

    elif imgtype == "objects":
        custom_prompt += """
        This image shows physical objects. Describe:
        - Type of object(s) (electronics, tools, toys, etc.)
        - Materials, shape, and condition
        - Background or context clues
        - Possible use or function of the objects
        """

    # Others — fallback
    elif imgtype == "others":
        custom_prompt += """
        Please:
        - Describe visible elements in detail
        - Mention key shapes, colors, patterns, and textures
        - Infer possible purpose or meaning
        - Keep it readable and structured
        """

    # Final summary generation
    response = model.generate_content([image, custom_prompt])
    session["summary"] = response.text
    return render_template("index.html", summary=response.text, image_url=image_display_path, mode="image")

def analyze_audio():
    audio = request.files.get("audio")
    if not audio or audio.filename == "":
        return render_template("index.html", summary="❌ No audio file uploaded."), 400

    filename = secure_filename(audio.filename)
    audio_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    audio.save(audio_path)

    if not os.path.exists(audio_path):
        return render_template("index.html", summary="❌ Audio file was not saved correctly."), 500

    try:
        logging.info("🧠 Transcribing audio with Whisper...")
        result = whisper_model.transcribe(audio_path)
        transcription = result["text"]
        logging.info(f"📜 Transcribed Text: {transcription}")
    except Exception as e:
        return render_template("index.html", summary=f"❌ Audio transcription failed: {str(e)}"), 500

    try:
        prompt = f"""
        The following is a transcription from an audio file. It may contain profanity or be in a language other than English.

        1. First, **remove any profane or offensive words** by replacing them with [redacted].
        2. Then, **translate the cleaned transcription into English**, if it's in another language.
        3. After cleaning and translating, analyze the content and do the following:

        a. Identify the type of audio:
            - casual conversation
            - formal discussion
            - song or music lyrics
            - speech or monologue
            - other (specify)

        b. Summarize the content in a few clean, structured lines:
            - If it’s a conversation, briefly explain what each person is talking about.
            - If it’s a song, describe the mood and the message.
            - If it’s noisy or unclear, mention that.

        4. Do not list or format dialogues. Just describe the intent or message behind them.
        5. Ensure the summary is clean, neutral, and free from offensive language.
        6. Keep everything concise, readable, and suitable for a general audience.
        7. Output only the final summary — no markdown, no labels, no symbols.

        Transcript:
        {transcription}
        """

        model = genai.GenerativeModel("models/gemini-2.0-flash")
        response = model.generate_content(prompt)
    except Exception as e:
        return render_template("index.html", summary=f"❌ Summarization failed: {str(e)}"), 500

    # ✅ Generate local audio file path for preview
    audio_url = url_for('static', filename=f'uploads/{filename}')
    session["summary"] = response.text
    return render_template("index.html", summary=response.text, audio_url=audio_url, mode="audio")


# Extract frames from video
def extract_frames(video_path, interval=90):
    frames = []
    cap = cv2.VideoCapture(video_path)
    frame_count = 0
    success, frame = cap.read()
    
    while success:
        if frame_count % interval == 0:
            frame_path = f"frame_{frame_count}.jpg"
            cv2.imwrite(frame_path, frame)
            frames.append(frame_path)
        success, frame = cap.read()
        frame_count += 1
    
    cap.release()
    return frames

# Describe frame using Gemini
def describe_frame(path):
    with Image.open(path) as image:  # ensures image is closed after use
        model = genai.GenerativeModel("models/gemini-2.0-flash")
        prompt = "Describe what's happening in this image."

        for attempt in range(5):
            try:
                response = model.generate_content([prompt, image])
                return response.text
            except google.api_core.exceptions.ResourceExhausted:
                logging.error(f"⚠️ Rate limit hit for image {path}. Retrying in 30 seconds...")
                time.sleep(30)
            except Exception as e:
                logging.error(f"❌ Unexpected error for image {path}: {e}")
                return f"[Error describing frame: {str(e)}]"

    return "[Failed to describe frame after multiple attempts due to rate limits.]"

# Extract audio and transcribe using Whisper
def extract_audio_and_transcribe(video_path):
    try:
        logging.info("🎵 Extracting audio...")
        video = mp.VideoFileClip(video_path)
        
        # Check if the video has audio
        if video.audio is None:
            logging.warning("⚠️ No audio stream found in the video.")
            return ""

        audio_path = "temp_audio.mp3"
        video.audio.write_audiofile(audio_path, logger=None)

        logging.info("🔤 Transcribing with Whisper...")
        model = whisper.load_model("base")  # Adjust model size as needed
        result = model.transcribe(audio_path)

        os.remove(audio_path)  # Clean up
        return result["text"]

    except Exception as e:
        logging.error(f"❌ Error during audio extraction or transcription: {e}")
        return ""

# Main pipeline
def analyse_video():
    video = request.files.get("video")
    if not video or video.filename == "":
        return render_template("index.html", summary="❌ No video file uploaded."), 400

    filename = secure_filename(video.filename)
    video_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    video.save(video_path)

    if not os.path.exists(video_path):
        return render_template("index.html", summary="❌ Video file was not saved correctly."), 500

    try:
        logging.info("📸 Extracting frames...")
        frame_paths = extract_frames(video_path)

        logging.info("🧠 Describing key frames...")
        descriptions = []
        for path in frame_paths:
            desc = describe_frame(path)
            logging.info(f"{path}: {desc}")
            descriptions.append(desc)
            os.remove(path)  # Clean up
    except Exception as e:
        return render_template("index.html", summary=f"❌ Frame analysis failed: {str(e)}"), 500

    try:
        transcript = extract_audio_and_transcribe(video_path)
    except Exception as e:
        return render_template("index.html", summary=f"❌ Audio transcription failed: {str(e)}"), 500

    try:
        prompt = f"""
        You are analyzing a video. Below are the extracted descriptions:

        Visual Descriptions:
        {chr(10).join(descriptions)}

        Audio Transcription:
        {transcript}

        Based on the combined visual and audio content, perform the following:

        1. Identify the main theme or context of this video (e.g., action scene, emotional dialogue, romantic clip, musical performance, news segment, etc.).

        2. Classify the video type:
        - Music video
        - Dialogue or conversation scene
        - Movie clip or trailer
        - Interview or speech
        - Other (specify)

        3. If possible, guess the name of the song, movie, or show featured in the video.

        4. Identify the language(s) spoken or sung in the video.

        5. Suggest any famous people, public figures, or actors who might be present, based on visuals or voices.

        6. Detect and mention the overall mood or emotional tone of the video (e.g., joyful, tense, dramatic, uplifting, sad, humorous).

        7. Detect and redact any profanity found in the transcription. Do not repeat offensive language — replace such words with [redacted].

        8. For each of your guesses (e.g., title, celebrities), include a confidence level (e.g., high, medium, low).

        9. If you can't answer any point, **just skip it** silently. Do not restate the question or say "unable to determine."

        10. Do not use markdown, bullet points, special formatting, or emojis.

        11. Your final answer should be clean, natural, and readable by anyone. Focus only on delivering insightful, structured observations from the video.
        """

        model = genai.GenerativeModel("models/gemini-2.0-flash")
        response = model.generate_content(prompt)
    except Exception as e:
        return render_template("index.html", summary=f"❌ Video summarization failed: {str(e)}"), 500

    video_url = url_for('static', filename=f'uploads/{filename}')
    session["summary"] = response.text
    return render_template("index.html", summary=response.text, video_url=video_url, mode="video")

@app.route('/download_summary')
def download_summary():
    summary = session.get("summary", "")
    if not summary:
        return "No summary available to download."

    doc = Document()
    doc.add_heading('AI Multimedia Analyzer Summary', level=1)
    doc.add_paragraph(summary)

    upload_folder = os.path.join(app.root_path, 'static', 'uploads')
    os.makedirs(upload_folder, exist_ok=True)
    doc_path = os.path.join(upload_folder, 'summary.docx')
    doc.save(doc_path)

    @after_this_request
    def remove_file(response):
        try:
            os.remove(doc_path)
        except Exception as e:
            print("Error deleting file:", e)
        return response

    return send_file(doc_path, as_attachment=True)


@app.errorhandler(413)
def too_large(e):
    return render_template("index.html", summary="❌ File too large. Max size 5MB."), 413

if __name__ == "__main__":
    app.run(debug=True)
